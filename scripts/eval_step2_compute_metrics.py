"""評測步驟 2：計算 Correctness / Faithfulness / Retrieval Precision / Response Time。

資料來源：
- 標準答案：C:\\Users\\user\\Desktop\\研究計畫\\AI詢問問題\\15題測試問題(修改)\\
- 360 份生成回答：C:\\Users\\user\\Desktop\\驗證\\第01~15題\\
- 檢索原文快取：C:\\Users\\user\\Desktop\\驗證\\數據資料\\retrieval_context_cache.json（步驟1產出）

公式出處（已對照官方論文/RAGAS 文件核實）：
- Correctness = 0.75*F1(事實重疊) + 0.25*SemanticSimilarity(餘弦相似度)  [RAGAS answer_correctness]
- Faithfulness = 被上下文支持的陳述句數 / 總陳述句數  [RAGAS faithfulness]
- Retrieval Precision = Context Precision@K  [RAGAS context_precision]

判斷 v_k（檢索結果是否相關）與 Correctness/Faithfulness 的事實拆解，
皆採用 LLM 裁判（qwen3:8b，溫度0）。
"""

import json
import re
import time
from pathlib import Path

import numpy as np
import requests
from docx import Document

OUTPUT_ROOT = Path("/mnt/c/Users/user/Desktop/驗證")
GROUND_TRUTH_DIR = Path("/mnt/c/Users/user/Desktop/研究計畫/AI詢問問題/15題測試問題(修改)")
DATA_DIR = OUTPUT_ROOT / "數據資料"
CONTEXT_CACHE_FILE = DATA_DIR / "retrieval_context_cache.json"
RESULTS_FILE = DATA_DIR / "metrics_results.json"
LOG_FILE = DATA_DIR / "eval_step2_log.txt"

OLLAMA_URL = "http://localhost:11434"
JUDGE_MODEL = "qwen3:8b"
EMBED_MODEL = "bge-m3:latest"

LLMS = [
    ("llama3.1:8b", "Llama3.1-8B"),
    ("deepseek-r1:8b", "DeepSeekR1-8B"),
    ("qwen3:8b", "Qwen3-8B"),
    ("deepseek-r1:32b", "DeepSeekR1-32B"),
    ("qwen3:30b", "Qwen3-30B"),
    ("gpt-oss:20b", "GPTOSS-20B"),
]
STRATEGIES = [
    ("naive_rag", "NaiveRAG"),
    ("advanced_rag", "AdvancedRAG"),
    ("graph_only", "GraphRAG"),
    ("agentic_rag", "AgenticRAG"),
]


def log(msg: str) -> None:
    line = f"[{time.strftime('%Y-%m-%d %H:%M:%S')}] {msg}"
    print(line, flush=True)
    with open(LOG_FILE, "a", encoding="utf-8") as f:
        f.write(line + "\n")


# ==================== 資料載入 ====================


def load_ground_truth() -> dict:
    """讀取 15 篇標準答案全文"""
    gt = {}
    for n in range(1, 16):
        path = GROUND_TRUTH_DIR / f"問題回答({n})(修改版).docx"
        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        gt[n - 1] = text
    return gt


def parse_generated_docx(path: Path) -> dict:
    """解析單一份生成回答 docx，取出題目、回答、來源、耗時"""
    doc = Document(str(path))
    paras = doc.paragraphs
    result = {"question": "", "answer": "", "sources": [], "elapsed": 0.0}
    section = None
    answer_lines = []
    for p in paras:
        text = p.text.strip()
        if not text:
            continue
        if text == "題目":
            section = "question"
            continue
        if text == "回答":
            section = "answer"
            continue
        if text == "參考來源":
            section = "sources"
            continue
        if text == "執行資訊":
            section = "info"
            continue
        if section == "question":
            result["question"] = text
        elif section == "answer":
            answer_lines.append(text)
        elif section == "sources" and text.startswith("-"):
            result["sources"].append(text)
        elif section == "info" and text.startswith("耗時"):
            m = re.search(r"([\d.]+)", text)
            if m:
                result["elapsed"] = float(m.group(1))
    result["answer"] = "\n".join(answer_lines)
    return result


def load_all_generated() -> dict:
    """回傳 {(q_idx, llm_tag, strategy_code): parsed_dict}"""
    out = {}
    for q_idx in range(15):
        folder = OUTPUT_ROOT / f"第{q_idx + 1:02d}題"
        for llm_tag, llm_name in LLMS:
            for strategy_code, rag_name in STRATEGIES:
                path = folder / f"{llm_name}_{rag_name}.docx"
                if not path.exists():
                    continue
                out[(q_idx, llm_tag, strategy_code)] = parse_generated_docx(path)
    return out


# ==================== LLM 裁判 / Embedding 呼叫 ====================


def ollama_generate(prompt: str, model: str = JUDGE_MODEL, temperature: float = 0.0) -> str:
    resp = requests.post(
        f"{OLLAMA_URL}/api/generate",
        json={
            "model": model,
            "prompt": prompt,
            "stream": False,
            "think": False,
            # temperature/num_predict 要放進 options 才會生效（放最外層 Ollama 會忽略，
            # 沒有 num_predict 上限代表可能無限生成）；裁判只需要短短一段 JSON，800 token夠用
            "options": {"temperature": temperature, "num_predict": 800},
        },
        timeout=180,
    )
    resp.raise_for_status()
    return resp.json().get("response", "")


def extract_json(text: str) -> dict:
    """從模型輸出中找出最後一個合法的 JSON 物件（容忍多餘文字/思考標記）"""
    text = re.sub(r"<think>.*?</think>", "", text, flags=re.DOTALL)
    matches = re.findall(r"\{[^{}]*\}", text, flags=re.DOTALL)
    for candidate in reversed(matches):
        try:
            return json.loads(candidate)
        except json.JSONDecodeError:
            continue
    return {}


def embed_text(text: str) -> np.ndarray:
    """呼叫 bge-m3 embedding；字數不等於 token 數，中文技術文本 token 密度可能偏高，
    因此若截斷長度仍超出模型上下文，逐步縮短重試。"""
    for limit in (2500, 1500, 800, 400):
        resp = requests.post(
            f"{OLLAMA_URL}/api/embed",
            json={"model": EMBED_MODEL, "input": text[:limit]},
            timeout=60,
        )
        if resp.status_code == 200:
            vec = resp.json()["embeddings"][0]
            return np.array(vec)
    resp.raise_for_status()
    vec = resp.json()["embeddings"][0]
    return np.array(vec)


def cosine_sim(a: np.ndarray, b: np.ndarray) -> float:
    denom = np.linalg.norm(a) * np.linalg.norm(b)
    if denom == 0:
        return 0.0
    return float(np.dot(a, b) / denom)


# ==================== 指標計算 ====================


def compute_correctness(question: str, reference: str, generated: str) -> dict:
    prompt = f"""你是嚴謹的事實查核員。請比較「待評估答案」與「標準答案」在事實內容上的重疊程度。

題目：{question}

標準答案：
{reference[:3000]}

待評估答案：
{generated[:3000]}

請執行以下步驟：
1. 將「待評估答案」拆解成獨立的事實陳述句
2. 對每一條陳述句判斷：
   - TP：此陳述句正確，且與標準答案內容一致或被其支持
   - FP：此陳述句與標準答案不一致，或標準答案未提及且無法確認為正確
3. 找出「標準答案」中有、但「待評估答案」完全沒提到的重要事實，計為 FN

只輸出以下 JSON，不要有其他文字：
{{"tp": 數字, "fp": 數字, "fn": 數字}}"""
    raw = ollama_generate(prompt)
    parsed = extract_json(raw)
    tp = parsed.get("tp", 0)
    fp = parsed.get("fp", 0)
    fn = parsed.get("fn", 0)
    denom = tp + 0.5 * (fp + fn)
    f1 = tp / denom if denom > 0 else 0.0

    try:
        e_gen = embed_text(generated)
        e_ref = embed_text(reference)
        sem_sim = cosine_sim(e_gen, e_ref)
        sem_sim = max(0.0, sem_sim)
    except Exception as e:
        log(f"    embedding 失敗: {e}")
        sem_sim = 0.0

    correctness = 0.75 * f1 + 0.25 * sem_sim
    return {
        "tp": tp,
        "fp": fp,
        "fn": fn,
        "f1": f1,
        "semantic_similarity": sem_sim,
        "correctness": correctness,
    }


def compute_faithfulness(generated: str, context_text: str) -> dict:
    if not context_text.strip():
        return {"total_claims": 0, "supported_claims": 0, "faithfulness": 0.0}
    prompt = f"""你是嚴謹的事實查核員。
請判斷「待評估答案」裡的每一條陳述句，是否能被「檢索到的參考資料」直接支持。

檢索到的參考資料：
{context_text[:4000]}

待評估答案：
{generated[:3000]}

請執行以下步驟：
1. 將「待評估答案」拆解成獨立的事實陳述句
2. 對每一條陳述句判斷是否能從上面的參考資料中找到直接支持的依據

只輸出以下 JSON，不要有其他文字：
{{"total_claims": 數字, "supported_claims": 數字}}"""
    raw = ollama_generate(prompt)
    parsed = extract_json(raw)
    total = parsed.get("total_claims", 0)
    supported = parsed.get("supported_claims", 0)
    faithfulness = supported / total if total > 0 else 0.0
    return {"total_claims": total, "supported_claims": supported, "faithfulness": faithfulness}


def compute_retrieval_precision(question: str, reference: str, sources: list) -> dict:
    if not sources:
        return {"relevance": [], "precision": 0.0}
    listing = "\n".join(
        f"[{i + 1}] 標題：{s.get('metadata', {}).get('title', '未知')}\n"
        f"內容：{(s.get('content') or '')[:300]}"
        for i, s in enumerate(sources)
    )
    prompt = f"""你是檢索相關性評估員。

題目：{question}
標準答案內容（作為相關性判斷基準）：
{reference[:2500]}

以下是檢索系統回傳的候選資料，請針對每一筆，判斷它是否跟這道題目「相關」——
也就是這筆資料的內容有沒有可能被用來回答或支持這道題目。

{listing}

只輸出以下 JSON（relevance 陣列長度需等於候選資料筆數，1=相關，0=不相關），不要有其他文字：
{{"relevance": [1或0, ...]}}"""
    raw = ollama_generate(prompt)
    parsed = extract_json(raw)
    relevance = parsed.get("relevance", [])
    if len(relevance) != len(sources):
        relevance = (relevance + [0] * len(sources))[: len(sources)]

    # Context Precision@K
    num_relevant = sum(relevance)
    if num_relevant == 0:
        precision = 0.0
    else:
        running_tp = 0
        precision_sum = 0.0
        for k, v_k in enumerate(relevance, start=1):
            if v_k:
                running_tp += 1
                precision_at_k = running_tp / k
                precision_sum += precision_at_k * v_k
        precision = precision_sum / num_relevant
    return {"relevance": relevance, "precision": precision}


# ==================== 主流程 ====================


def main():
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    log("載入標準答案...")
    ground_truth = load_ground_truth()
    log("載入 360 份生成回答...")
    generated_all = load_all_generated()
    log(f"共 {len(generated_all)} 份")
    context_cache = json.loads(CONTEXT_CACHE_FILE.read_text(encoding="utf-8"))

    results = {}
    if RESULTS_FILE.exists():
        results = json.loads(RESULTS_FILE.read_text(encoding="utf-8"))

    retrieval_precision_cache = {}  # key: "q_idx|strategy" -> result（跨 6 個 LLM 共用）

    total = len(generated_all)
    count = 0
    for (q_idx, llm_tag, strategy_code), gen in generated_all.items():
        count += 1
        key = f"{q_idx}|{llm_tag}|{strategy_code}"
        if key in results:
            continue

        log(f"[{count}/{total}] 第{q_idx + 1}題 - {llm_tag} + {strategy_code} 計算中...")
        reference = ground_truth[q_idx]
        question = gen["question"]
        answer = gen["answer"]

        ctx_key = f"{q_idx}|{strategy_code}"
        ctx_sources = context_cache.get(ctx_key, [])
        context_text = "\n\n".join((s.get("content") or "") for s in ctx_sources)

        try:
            correctness_result = compute_correctness(question, reference, answer)
        except Exception as e:
            log(f"  Correctness 計算失敗: {e}")
            correctness_result = {"correctness": 0.0}

        try:
            faithfulness_result = compute_faithfulness(answer, context_text)
        except Exception as e:
            log(f"  Faithfulness 計算失敗: {e}")
            faithfulness_result = {"faithfulness": 0.0}

        rp_key = ctx_key
        if rp_key not in retrieval_precision_cache:
            try:
                retrieval_precision_cache[rp_key] = compute_retrieval_precision(
                    question, reference, ctx_sources
                )
            except Exception as e:
                log(f"  Retrieval Precision 計算失敗: {e}")
                retrieval_precision_cache[rp_key] = {"precision": 0.0}
        rp_result = retrieval_precision_cache[rp_key]

        results[key] = {
            "correctness": correctness_result["correctness"],
            "faithfulness": faithfulness_result["faithfulness"],
            "retrieval_precision": rp_result["precision"],
            "response_time": gen["elapsed"],
        }
        RESULTS_FILE.write_text(json.dumps(results, ensure_ascii=False, indent=1), encoding="utf-8")

    log(f"=== 全部完成，共 {len(results)} 筆 ===")


if __name__ == "__main__":
    main()
