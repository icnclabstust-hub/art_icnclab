"""24 組合（6 LLM x 4 RAG）x 15 題評測，每題每組合輸出一份 .docx。

輸出結構：C:\\Users\\user\\Desktop\\驗證\\第XX題\\LLM名稱_RAG名稱.docx
支援中斷續跑：progress.json 記錄已完成的 (question_idx, llm_key, strategy_key)。
"""

import json
import time
from datetime import datetime
from pathlib import Path

import requests
from docx import Document

RAG_MANAGER_URL = "http://localhost:8007"
OUTPUT_ROOT = Path("/mnt/c/Users/user/Desktop/驗證")
PROGRESS_FILE = OUTPUT_ROOT / "progress.json"
LOG_FILE = OUTPUT_ROOT / "run_log.txt"

QUESTIONS = [
    "辨識「Opera」「Opera del Duomo」「Opera di Santa Maria del Fiore」"
    "在混合語料資料庫中的指涉實體。",
    "Figura 的英語術語對應、歷史—神學義及其在教父至早期文藝復興中的承接。",
    "Disegno 的歷史收縮。",
    "Kunstwollen 的四層結構。",
    "達文西的邊界理論與暈塗法的光學優位。",
    "Medici Bank、Poggio 網絡與聖羅倫佐教堂：十五世紀佛羅倫斯秩序形式的比較考察。",
    "十四世紀煉金蒸餾、尼德蘭媒介改良與北方光學寫實。",
    "米開朗基羅天頂畫中先知與巫女的空間重構。",
    "從威尼斯 Colorito 到魯本斯祭壇畫色彩邏輯。",
    "Isabella d'Este 的怪誕審美與瓦薩里式藝術史的張力。",
    "《雅典學院》與《夜巡》的形式史深度對比。",
    "拉斐爾《雅典學院》中柏拉圖與亞里斯多德手勢的哲學分歧。",
    "從材料化學與光學散射看盛期文藝復興濕壁畫與威尼斯油畫的內部發光感。",
    "從「威壓的生命能量」到 Terribilità。",
    "伊兒汗細密畫、地中海節點與特雷琴托祭壇畫。",
]

LLMS = [
    ("llama3.1:8b", "Llama3.1-8B"),
    ("deepseek-r1:8b", "DeepSeekR1-8B"),
    ("qwen3:8b", "Qwen3-8B"),
    ("deepseek-r1:32b", "DeepSeekR1-32B"),
    ("qwen3:30b", "Qwen3-30B"),
    ("gpt-oss:20b", "GPTOSS-20B"),
]

STRATEGIES = [
    ("naive_rag", "NaiveRAG"),
    ("advanced_rag", "AdvancedRAG"),
    ("graph_only", "GraphRAG"),
    ("agentic_rag", "AgenticRAG"),
]


def log(msg: str) -> None:
    line = f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] {msg}"
    print(line, flush=True)
    with open(LOG_FILE, "a", encoding="utf-8") as f:
        f.write(line + "\n")


def load_progress() -> set:
    if PROGRESS_FILE.exists():
        return set(tuple(x) for x in json.loads(PROGRESS_FILE.read_text(encoding="utf-8")))
    return set()


def save_progress(done: set) -> None:
    PROGRESS_FILE.write_text(json.dumps(sorted(list(done)), ensure_ascii=False), encoding="utf-8")


def query_rag(question: str, llm_tag: str, strategy_code: str):
    combo_id = f"{llm_tag}@{strategy_code}"
    start = time.time()
    resp = requests.post(
        f"{RAG_MANAGER_URL}/api/v1/query",
        json={
            "query": question,
            "model_combination_id": combo_id,
            "max_results": 5,
            "include_sources": True,
        },
        timeout=950,  # 需大於伺服端 Ollama 生成逾時(900s)，保留餘裕給檢索/前處理步驟
    )
    elapsed = time.time() - start
    resp.raise_for_status()
    data = resp.json()
    return data, elapsed


def write_docx(
    path: Path, q_idx: int, question: str, llm_name: str, rag_name: str, data, elapsed, error=None
):
    doc = Document()
    doc.add_heading(f"第{q_idx + 1:02d}題 - {llm_name} + {rag_name}", level=1)

    doc.add_heading("題目", level=2)
    doc.add_paragraph(question)

    doc.add_heading("回答", level=2)
    if error:
        doc.add_paragraph(f"[錯誤] {error}")
    else:
        doc.add_paragraph(data.get("answer", "(無回答)"))

        sources = data.get("sources", [])
        if sources:
            doc.add_heading("參考來源", level=2)
            for s in sources:
                meta = s.get("metadata", {})
                title = meta.get("title", "未知標題")
                source = meta.get("source", "未知來源")
                score = s.get("score", "")
                doc.add_paragraph(f"- {title}  [{source}]  score={score}", style="List Bullet")

    doc.add_heading("執行資訊", level=2)
    doc.add_paragraph(f"LLM+RAG 組合: {llm_name} + {rag_name}")
    doc.add_paragraph(f"耗時: {elapsed:.1f} 秒")
    doc.add_paragraph(f"執行時間: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def main():
    OUTPUT_ROOT.mkdir(parents=True, exist_ok=True)
    done = load_progress()
    total = len(LLMS) * len(STRATEGIES) * len(QUESTIONS)
    log(f"=== 開始 24 組合 x 15 題評測，共 {total} 次查詢，已完成 {len(done)} 次 ===")

    count = len(done)
    for llm_tag, llm_name in LLMS:
        for strategy_code, rag_name in STRATEGIES:
            for q_idx, question in enumerate(QUESTIONS):
                key = (q_idx, llm_tag, strategy_code)
                key_str = list(key)
                if tuple(key_str) in done:
                    continue

                folder = OUTPUT_ROOT / f"第{q_idx + 1:02d}題"
                filename = f"{llm_name}_{rag_name}.docx"
                filepath = folder / filename

                log(f"[{count + 1}/{total}] 第{q_idx + 1}題 - {llm_name} + {rag_name} ...")
                try:
                    data, elapsed = query_rag(question, llm_tag, strategy_code)
                    write_docx(filepath, q_idx, question, llm_name, rag_name, data, elapsed)
                    log(f"  ✅ 完成，耗時 {elapsed:.1f} 秒")
                except Exception as e:
                    log(f"  ❌ 失敗: {e}")
                    write_docx(filepath, q_idx, question, llm_name, rag_name, None, 0, error=str(e))

                done.add(key)
                save_progress(done)
                count += 1

    log(f"=== 全部完成，共 {count} 次查詢 ===")

    manifest_path = OUTPUT_ROOT / "manifest.txt"
    all_files = sorted(OUTPUT_ROOT.glob("第*題/*.docx"))
    with open(manifest_path, "w", encoding="utf-8") as f:
        f.write(f"共 {len(all_files)} 個檔案\n\n")
        for fp in all_files:
            f.write(str(fp.relative_to(OUTPUT_ROOT)) + "\n")
    log(f"清單已寫入 {manifest_path}")


if __name__ == "__main__":
    main()
